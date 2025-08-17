"""
Loads and embeds documents into a FAISS vectorstore for retrieval.
"""

import json
import time
import hashlib
import pandas as pd
from datetime import datetime
from hashlib import md5
from pathlib import Path
from typing import List, Dict, Optional
from app.utils import (
    log_info,
    log_success,
    log_warning,
    log_error,
    timer,
    validate_path,
    create_progress_bar,
)
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document
from app.chunking import SemanticChunker

MANIFEST_NAME = "manifest.json"
EMBEDDING_MODEL = "nomic-embed-text"
SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".txt",
    ".csv",
    ".md",
    ".rtf",
    # Images (for future OCR/vision)
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".bmp",
    ".tiff",
    # Code files
    ".py",
    ".js",
    ".ts",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".cs",
    ".go",
    ".rs",
    ".php",
    ".html",
    ".css",
    ".xml",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".sql",
    ".sh",
    ".bat",
    ".ps1",
    ".r",
    ".m",
    ".swift",
    ".kt",
    ".dart",
    # Text/Documentation
    ".org",
    ".rst",
    ".tex",
    ".log",
    ".conf",
    ".properties",
    # Email (common formats)
    ".eml",
    ".msg",
    ".mbox",
    # Archives (we can extract and process contents)
    ".zip",
    ".tar",
    ".gz",
    # Bookmarks
    ".html",  # Browser bookmarks export as HTML
}


def file_checksum(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while chunk := f.read(8192):
            h.update(chunk)
    return h.hexdigest()


def file_md5(path: Path, buf_size: int = 8 * 1024 * 1024) -> str:
    h = md5()
    with path.open("rb") as f:
        while True:
            chunk = f.read(buf_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def load_manifest(persist_path: str) -> Dict[str, str]:
    p = Path(persist_path) / MANIFEST_NAME
    if not p.exists():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_manifest(persist_path: str, mapping: Dict[str, str]):
    p = Path(persist_path) / MANIFEST_NAME
    p.write_text(json.dumps(mapping, ensure_ascii=False, indent=2), encoding="utf-8")


def chunk_documents(texts: List[Dict], chunk_size: int, chunk_overlap: int):
    """Enhanced document chunking with semantic awareness."""
    chunker = SemanticChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    all_chunks = []

    for doc in texts:
        try:
            # Use semantic chunking
            chunks = chunker.chunk_document(doc["text"], doc["metadata"])
            all_chunks.extend(chunks)
        except Exception as e:
            log_warning(
                f"Semantic chunking failed for {doc['metadata'].get('source', 'unknown')}: {e}"
            )
            # Fallback to basic chunking
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            basic_chunks = splitter.split_text(doc["text"])
            for i, chunk in enumerate(basic_chunks):
                all_chunks.append(
                    {
                        "text": chunk,
                        "metadata": {
                            **doc["metadata"],
                            "chunk_id": f"{doc['metadata']['source']}#{i}",
                            "chunk_type": "fallback",
                        },
                    }
                )

    return all_chunks


def normalize_documents(docs: list[Document]) -> list[Document]:
    """
    Normalizes document content and metadata for consistency.

    Args:
        docs (list[Document]): List of documents to normalize.

    Returns:
        list[Document]: List of normalized documents.
    """
    norm = []
    for d in docs:
        content = (d.page_content or "").strip()
        if not content:
            continue
        meta = dict(d.metadata or {})
        meta.setdefault("source", meta.get("file_path", "unknown"))
        norm.append(Document(page_content=content, metadata=meta))
    return norm


def dedupe_documents_by_content(docs: list[Document]) -> list[Document]:
    """
    Removes duplicate documents based on their content hash.

    Args:
        docs (list[Document]): List of documents to deduplicate.

    Returns:
        list[Document]: List of deduplicated documents.
    """
    seen = set()
    out = []
    for d in docs:
        h = md5(d.page_content.encode("utf-8")).hexdigest()
        if h in seen:
            continue
        seen.add(h)
        out.append(d)
    return out


def write_index_meta(
    persist_path: str, *, embedding_model: str, chunk_size: int, chunk_overlap: int
):
    """
    Writes metadata for the FAISS index to disk.

    Args:
        persist_path (str): Path to the directory where metadata will be saved.
        embedding_model (str): Name of the embedding model used.
        chunk_size (int): Size of the document chunks.
        chunk_overlap (int): Overlap between document chunks.

    Returns:
        None
    """
    meta = {
        "created_at": time.time(),
        "embedding_model": embedding_model,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
    }
    with open(Path(persist_path) / "metadata.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def read_index_meta(persist_path: str) -> Optional[dict]:
    """
    Reads metadata for the FAISS index from disk.

    Args:
        persist_path (str): Path to the directory where metadata is stored.

    Returns:
        dict or None: Dictionary containing the metadata, or None if not found.
    """
    p = Path(persist_path) / "metadata.json"
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


def load_excel_with_pandas(file_path: Path) -> list[Document]:
    """
    Pure-Python Excel loader using pandas; returns one Document per sheet.

    Args:
        file_path: Path to the Excel file to load

    Returns:
        List of Document objects, one per sheet in the Excel file
    """
    docs = []
    try:
        sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
        for sheet_name, df in sheets.items():
            if df.empty:
                continue
            text = df.to_csv(index=False)
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": str(file_path),
                        "sheet": sheet_name,
                        "type": "excel",
                        "filename": file_path.name,
                        "ext": file_path.suffix.lower(),
                    },
                )
            )
    except Exception as e:
        log_error(f"Excel load failed for {file_path.name}: {e}")
    return docs


def extract_metadata(file_path: Path, file_type: str, extra=None):
    meta = {
        "source": str(file_path),
        "type": file_type,
        "filename": file_path.name,
        "ext": file_path.suffix.lower(),
        "hash": (
            md5(file_path.read_bytes()).hexdigest() if file_path.is_file() else None
        ),
    }
    try:
        if file_type == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            meta["title"] = getattr(reader.metadata, "title", None)
            meta["author"] = getattr(reader.metadata, "author", None)
        elif file_type == ".docx":
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            props = doc.core_properties
            meta["author"] = props.author
            meta["title"] = props.title
        elif file_type in {".xlsx", ".xls"}:
            import openpyxl

            wb = openpyxl.load_workbook(str(file_path), read_only=True)
            meta["sheets"] = wb.sheetnames
            meta["creator"] = wb.properties.creator
        elif file_type == ".txt":
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
                meta["title"] = lines[0].strip() if lines else None
        elif file_type in {".png", ".jpg", ".jpeg", ".webp"}:
            from PIL import Image

            with Image.open(str(file_path)) as img:
                meta["image_size"] = img.size
                meta["image_mode"] = img.mode
        elif file_type in {".py", ".c", ".cpp", ".html"}:
            meta["lines"] = sum(
                1 for _ in open(file_path, encoding="utf-8", errors="ignore")
            )
    except Exception as e:
        meta["meta_error"] = str(e)
    if extra:
        meta.update(extra)
    return meta


def get_loader_for_file(file_path: Path):
    """
    Returns the appropriate document loader for a given file type.

    Args:
        file_path (Path): Path to the file.

    Returns:
        Document loader instance or None if unsupported.
    """
    suffix = file_path.suffix.lower()

    # Document formats
    if suffix == ".pdf":
        return PyPDFLoader(str(file_path))
    if suffix == ".docx":
        return Docx2txtLoader(str(file_path))
    if suffix in {".txt", ".md", ".rst", ".org", ".log", ".conf", ".properties"}:
        return TextLoader(str(file_path), autodetect_encoding=True)

    # Code files - treat as text
    if suffix in {
        ".py",
        ".js",
        ".ts",
        ".java",
        ".c",
        ".cpp",
        ".h",
        ".hpp",
        ".cs",
        ".go",
        ".rs",
        ".php",
        ".html",
        ".css",
        ".xml",
        ".json",
        ".yaml",
        ".yml",
        ".toml",
        ".ini",
        ".cfg",
        ".sql",
        ".sh",
        ".bat",
        ".ps1",
        ".r",
        ".m",
        ".swift",
        ".kt",
        ".dart",
    }:
        return TextLoader(str(file_path), autodetect_encoding=True)

    # Excel files - handled separately
    if suffix in {".xlsx", ".xls"}:
        return None  # Special handling in load_documents

    # Images - metadata only for now (future: OCR/vision)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}:
        return "metadata_only"

    # Email formats - need special handling
    if suffix in {".eml", ".msg", ".mbox"}:
        return "email_loader"  # Will implement this

    # Archives - need extraction
    if suffix in {".zip", ".tar", ".gz"}:
        return "archive_loader"  # Will implement this

    return None


def load_documents(folder_path: str) -> List[Document]:
    """
    Loads documents from the given folder with improved error handling and progress tracking.
    Supported formats: PDF, DOCX, Excel, TXT.

    Args:
        folder_path: Path to folder containing documents

    Returns:
        List of loaded documents
    """
    try:
        folder = validate_path(folder_path, must_exist=True)
    except (FileNotFoundError, NotADirectoryError) as e:
        log_error(str(e))
        return []

    supported_files = [
        f
        for f in folder.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    if not supported_files:
        log_warning(f"No supported files found in '{folder_path}'")
        return []

    docs, failed_files, ingest_events = [], [], []

    with create_progress_bar("Loading documents") as progress:
        task = progress.add_task("Loading...", total=len(supported_files))
        for file_path in supported_files:
            suffix = file_path.suffix.lower()
            try:
                meta = extract_metadata(file_path, suffix)
                if suffix in {".xlsx", ".xls"}:
                    file_docs = load_excel_with_pandas(file_path)
                    for doc in file_docs:
                        doc.metadata.update(meta)
                    docs.extend(file_docs)
                    ingest_events.append(
                        {"file": file_path.name, "status": "success", "meta": meta}
                    )
                else:
                    loader = get_loader_for_file(file_path)
                    if loader == "metadata_only":
                        docs.append(Document(page_content="", metadata=meta))
                        ingest_events.append(
                            {
                                "file": file_path.name,
                                "status": "metadata_only",
                                "meta": meta,
                            }
                        )
                    elif loader:
                        file_docs = loader.load()
                        for doc in file_docs:
                            doc.metadata.update(meta)
                        docs.extend(file_docs)
                        ingest_events.append(
                            {"file": file_path.name, "status": "success", "meta": meta}
                        )
                    else:
                        ingest_events.append(
                            {
                                "file": file_path.name,
                                "status": "no_loader",
                                "meta": meta,
                            }
                        )
            except Exception as e:
                failed_files.append(file_path.name)
                ingest_events.append(
                    {"file": file_path.name, "status": "failed", "error": str(e)}
                )
            progress.advance(task)

    log_dir = Path("data/ingest")
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"{datetime.now().strftime('%m-%d-%Y_%H-%M-%S')}.log"
    with open(log_file, "w", encoding="utf-8") as f:
        for event in ingest_events:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")

    if failed_files:
        log_warning(
            f"Failed to load {len(failed_files)} files: {', '.join(failed_files)}"
        )
    log_success(
        f"Loaded {len(docs)} document chunks from {len(supported_files) - len(failed_files)} files"
    )
    return docs


@timer
def build_vectorstore(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
    mode: str = "rebuild",  # or "increment"
):
    """
    Builds or updates a FAISS vectorstore from documents in a folder.

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.
        mode (str): 'rebuild' to wipe and rebuild, 'increment' to update existing index.

    Returns:
        FAISS vectorstore or None if failed.
    """
    if mode == "rebuild":
        return rebuild_vectorstore(
            folder_path, persist_path, chunk_size, chunk_overlap, embedding_model
        )
    elif mode == "increment":
        return incremental_vectorstore(
            folder_path, persist_path, chunk_size, chunk_overlap, embedding_model
        )
    else:
        log_error(f"Invalid mode '{mode}'. Use 'rebuild' or 'increment'.")
        return None


@timer
def rebuild_vectorstore(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
) -> Optional[FAISS]:
    """
    Wipes and rebuilds the FAISS vectorstore from all documents in the folder.

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.

    Returns:
        Optional[FAISS]: The rebuilt FAISS vectorstore, or None if failed.
    """
    folder = Path(folder_path)
    manifest = {"files": {}}
    docs = []

    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        loader = get_loader_for_file(path)
        if loader == "metadata_only":
            text = ""
        elif loader:
            text = "\n".join([d.page_content for d in loader.load()])
        else:
            continue
        if text.strip() or loader == "metadata_only":
            docs.append({"text": text, "metadata": {"source": str(path)}})
            manifest["files"][str(path)] = file_checksum(path)

    chunks = chunk_documents(docs, chunk_size, chunk_overlap)
    texts = [c["text"] for c in chunks]
    metadatas = [c["metadata"] for c in chunks]

    embeddings = OllamaEmbeddings(model=embedding_model)
    vectorstore = FAISS.from_texts(texts, embedding=embeddings, metadatas=metadatas)
    vectorstore.save_local(persist_path)
    save_manifest(persist_path, manifest)

    log_success(
        f"Rebuilt vectorstore with {len(chunks)} chunks from {len(docs)} documents."
    )
    return True


@timer
def incremental_vectorstore(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
) -> Optional[FAISS]:
    """
    Incrementally updates an existing FAISS index:
    - Adds new/changed documents
    - Removes deleted documents

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.

    Returns:
        Optional[FAISS]: The updated FAISS vectorstore, or None if failed.
    """
    folder = Path(folder_path)
    old_manifest = load_manifest(persist_path)
    new_manifest = {"files": {}}
    changed_docs = []

    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        checksum = file_checksum(path)
        new_manifest["files"][str(path)] = checksum
        if old_manifest["files"].get(str(path)) != checksum:
            loader = get_loader_for_file(path)
            if loader == "metadata_only":
                text = ""
            elif loader:
                text = "\n".join([d.page_content for d in loader.load()])
            else:
                continue
            if text.strip() or loader == "metadata_only":
                changed_docs.append({"text": text, "metadata": {"source": str(path)}})

    removed = set(old_manifest["files"]) - set(new_manifest["files"])
    if removed:
        log_info(f"Detected {len(removed)} deleted files: {removed}")

    if changed_docs or removed:
        return rebuild_vectorstore(
            folder_path, persist_path, chunk_size, chunk_overlap, embedding_model
        )

    log_info("No document changes detected. Vectorstore is up-to-date.")
    return True
