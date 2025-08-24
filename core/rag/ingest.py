"""
Loads and embeds documents into a FAISS vectorstore for retrieval.
Enhanced with async processing and smart caching for better performance.
"""

import json
import time
import hashlib
import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional
from core.utils.orion_utils import (
    log_info,
    log_success,
    log_warning,
    log_error,
    timer,
    validate_path,
    create_progress_bar,
)
from core.processing.async_processing import AsyncDocumentProcessor
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document
from core.processing.chunking import SemanticChunker
from core.processing.document_intelligence import (
    MetadataEnricher,
    SmartChunker,
    DocumentMetadata,
)

MANIFEST_NAME = "manifest.json"
EMBEDDING_MODEL = "nomic-embed-text"
SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".txt",
    ".csv",
    ".md",
    ".rtf",
    # Images (for future OCR/vision)
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".bmp",
    ".tiff",
    # Code files
    ".py",
    ".js",
    ".ts",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".cs",
    ".go",
    ".rs",
    ".php",
    ".html",
    ".css",
    ".xml",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".sql",
    ".sh",
    ".bat",
    ".ps1",
    ".r",
    ".m",
    ".swift",
    ".kt",
    ".dart",
    # Text/Documentation
    ".org",
    ".rst",
    ".tex",
    ".log",
    ".conf",
    ".properties",
    # Email (common formats)
    ".eml",
    ".msg",
    ".mbox",
    # Archives (we can extract and process contents)
    ".zip",
    ".tar",
    ".gz",
    ".7z",
    # Bookmarks
    ".html",
}


def file_checksum(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while chunk := f.read(8192):
            h.update(chunk)
    return h.hexdigest()


def file_md5(path: Path, buf_size: int = 8 * 1024 * 1024) -> str:
    h = hashlib.md5()
    with path.open("rb") as f:
        while True:
            chunk = f.read(buf_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def load_manifest(persist_path: str) -> Dict[str, str]:
    p = Path(persist_path) / MANIFEST_NAME
    if not p.exists():
        return {}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_manifest(persist_path: str, mapping: Dict[str, str]):
    p = Path(persist_path) / MANIFEST_NAME
    p.write_text(json.dumps(mapping, ensure_ascii=False, indent=2), encoding="utf-8")


def chunk_documents(texts: List[Dict], chunk_size: int, chunk_overlap: int):
    """Enhanced document chunking with document intelligence."""
    # Initialize our smart chunker
    smart_chunker = SmartChunker(chunk_size=chunk_size, overlap=chunk_overlap)

    # Fallback semantic chunker for edge cases
    semantic_chunker = SemanticChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    all_chunks = []

    for doc in texts:
        try:
            # Convert metadata dict to DocumentMetadata if needed
            if isinstance(doc["metadata"], dict):
                # Create a basic DocumentMetadata from the dict
                doc_metadata = DocumentMetadata(
                    filename=doc["metadata"].get("filename", "unknown"),
                    filepath=doc["metadata"].get("source", "unknown"),
                    file_size=len(doc["text"].encode("utf-8")),
                    created_date=datetime.now(),  # Use current time as fallback
                    modified_date=datetime.now(),  # Use current time as fallback
                    document_type=doc["metadata"].get("document_type", "text_plain"),
                    word_count=len(doc["text"].split()),
                    line_count=len(doc["text"].split("\n")),
                )
            else:
                doc_metadata = doc["metadata"]

            # Use smart chunking with document intelligence
            chunks = smart_chunker.chunk_document(doc["text"], doc_metadata)

            # Add intelligence metadata to chunks
            for chunk in chunks:
                # Preserve existing metadata and add chunking info
                chunk["metadata"].update(
                    {
                        "chunking_method": "smart",
                        "original_source": doc["metadata"].get("source"),
                    }
                )

            all_chunks.extend(chunks)
            log_info(f"Smart chunking created {len(chunks)} chunks for " f"{doc['metadata'].get('source', 'unknown')}")

        except Exception as e:
            log_warning(
                f"Smart chunking failed for {doc['metadata'].get('source', 'unknown')}: {e}"
                f" - Falling back to semantic chunking"
            )
            try:
                # Fallback to semantic chunking
                chunks = semantic_chunker.chunk_document(doc["text"], doc["metadata"])
                for chunk in chunks:
                    chunk["metadata"]["chunking_method"] = "semantic_fallback"
                all_chunks.extend(chunks)
            except Exception as e2:
                log_warning(
                    f"Semantic chunking also failed for {doc['metadata'].get('source', 'unknown')}: {e2}"
                    f" - Using basic chunking"
                )
                # Final fallback to basic chunking
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                )
                basic_chunks = splitter.split_text(doc["text"])
                for i, chunk in enumerate(basic_chunks):
                    all_chunks.append(
                        {
                            "text": chunk,
                            "metadata": {
                                **doc["metadata"],
                                "chunk_id": f"{doc['metadata']['source']}#{i}",
                                "chunk_type": "fallback",  # Keep test compatibility
                                "chunking_method": "basic",
                            },
                        }
                    )

    log_info(f"Total chunks created: {len(all_chunks)}")
    return all_chunks


def normalize_documents(docs: list[Document]) -> list[Document]:
    """
    Normalizes document content and metadata for consistency.

    Args:
        docs (list[Document]): List of documents to normalize.

    Returns:
        list[Document]: List of normalized documents.
    """
    norm = []
    for d in docs:
        content = (d.page_content or "").strip()
        if not content:
            continue
        meta = dict(d.metadata or {})
        meta.setdefault("source", meta.get("file_path", "unknown"))
        norm.append(Document(page_content=content, metadata=meta))
    return norm


def dedupe_documents_by_content(docs: list[Document]) -> list[Document]:
    """
    Removes duplicate documents based on their content hash.

    Args:
        docs (list[Document]): List of documents to deduplicate.

    Returns:
        list[Document]: List of deduplicated documents.
    """
    seen = set()
    out = []
    for d in docs:
        h = hashlib.md5(d.page_content.encode("utf-8")).hexdigest()
        if h in seen:
            continue
        seen.add(h)
        out.append(d)
    return out


def write_index_meta(persist_path: str, *, embedding_model: str, chunk_size: int, chunk_overlap: int):
    """
    Writes metadata for the FAISS index to disk.

    Args:
        persist_path (str): Path to the directory where metadata will be saved.
        embedding_model (str): Name of the embedding model used.
        chunk_size (int): Size of the document chunks.
        chunk_overlap (int): Overlap between document chunks.

    Returns:
        None
    """
    meta = {
        "created_at": time.time(),
        "embedding_model": embedding_model,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
    }
    with open(Path(persist_path) / "metadata.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def read_index_meta(persist_path: str) -> Optional[dict]:
    """
    Reads metadata for the FAISS index from disk.

    Args:
        persist_path (str): Path to the directory where metadata is stored.

    Returns:
        dict or None: Dictionary containing the metadata, or None if not found.
    """
    p = Path(persist_path) / "metadata.json"
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


def load_excel_with_pandas(file_path: Path) -> list[Document]:
    """
    Pure-Python Excel loader using pandas; returns one Document per sheet.

    Args:
        file_path: Path to the Excel file to load

    Returns:
        List of Document objects, one per sheet in the Excel file
    """
    docs = []
    try:
        sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
        for sheet_name, df in sheets.items():
            if df.empty:
                continue
            text = df.to_csv(index=False)
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": str(file_path),
                        "sheet": sheet_name,
                        "type": "excel",
                        "filename": file_path.name,
                        "ext": file_path.suffix.lower(),
                    },
                )
            )
    except Exception as e:
        log_error(f"Excel load failed for {file_path.name}: {e}")
    return docs


def extract_metadata(file_path: Path, file_type: str, extra=None):
    """Enhanced metadata extraction using document intelligence, media processing, and code analysis"""
    # Basic metadata (original functionality preserved)
    meta = {
        "source": str(file_path),
        "type": file_type,
        "filename": file_path.name,
        "ext": file_path.suffix.lower(),
        "hash": (hashlib.md5(file_path.read_bytes()).hexdigest() if file_path.is_file() else None),
    }

    try:
        # Use document intelligence for enhanced metadata
        if file_path.is_file():
            # Read file content for analysis
            content = ""
            try:
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    content = f.read()
            except Exception as e:
                log_warning(f"Could not read file content for intelligence analysis: {e}")

            if content:
                # Use our document intelligence system
                enricher = MetadataEnricher()
                doc_metadata = enricher.extract_metadata(str(file_path), content)

                # Merge enhanced metadata with basic metadata
                meta.update(
                    {
                        "document_type": doc_metadata.document_type,
                        "title": doc_metadata.title or meta.get("title"),
                        "author": doc_metadata.author,
                        "word_count": doc_metadata.word_count,
                        "line_count": doc_metadata.line_count,
                        "paragraph_count": doc_metadata.paragraph_count,
                        "topics": doc_metadata.topics,
                        "keywords": doc_metadata.keywords,
                        "programming_language": doc_metadata.programming_language,
                        "functions_detected": doc_metadata.functions_detected,
                        "classes_detected": doc_metadata.classes_detected,
                        "imports_detected": doc_metadata.imports_detected,
                        "content_hash": doc_metadata.content_hash,
                        "intelligence_confidence": doc_metadata.confidence_score,
                    }
                )

                log_info(
                    f"Enhanced metadata extracted: {doc_metadata.document_type} "
                    f"({doc_metadata.word_count} words, {len(doc_metadata.topics)} topics)"
                )

                # Phase 1 & 2: Enhanced processing for specific file types
                try:
                    # Image processing for OCR and analysis
                    image_extensions = {
                        ".png",
                        ".jpg",
                        ".jpeg",
                        ".webp",
                        ".gif",
                        ".bmp",
                        ".tiff",
                    }

                    if file_type in image_extensions:
                        from core.processing.media_processing import media_processor

                        image_analysis = media_processor.process_image(file_path)

                        if image_analysis.get("success"):
                            ocr_result = image_analysis.get("ocr", {})
                            meta.update(
                                {
                                    "image_analysis": {
                                        "dimensions": (
                                            f"{image_analysis['metadata']['width']}" f"x{image_analysis['metadata']['height']}"
                                        ),
                                        "format": image_analysis["metadata"]["format"],
                                        "mode": image_analysis["metadata"]["mode"],
                                        "ocr_text_length": len(ocr_result.get("text", "")),
                                        "ocr_confidence": ocr_result.get("confidence", 0),
                                        "ocr_method": ocr_result.get("method", "none"),
                                        "tables_found": len(image_analysis.get("tables", [])),
                                        "processing_time": image_analysis.get("processing_time", 0),
                                    }
                                }
                            )

                            # If OCR found text, add it to the document content
                            ocr_text = ocr_result.get("text", "").strip()
                            if ocr_text and len(ocr_text) > 10:  # Only if substantial text found
                                meta["ocr_text"] = ocr_text
                                log_info(f"OCR extracted {len(ocr_text)} characters from image")

                except Exception as e:
                    log_warning(f"Enhanced processing failed for {file_path}: {e}")

        # Enhanced PDF processing (Phase 2: Table detection)
        if file_type == ".pdf":
            try:
                from pypdf import PdfReader
                from core.processing.media_processing import media_processor

                reader = PdfReader(str(file_path))
                # Only override if document intelligence didn't find a title
                if not meta.get("title"):
                    meta["title"] = getattr(reader.metadata, "title", None)
                if not meta.get("author"):
                    meta["author"] = getattr(reader.metadata, "author", None)

                # Phase 2: Enhanced PDF processing with table detection
                pdf_analysis = media_processor.process_pdf_enhanced(file_path)
                if pdf_analysis.get("success"):
                    tables = pdf_analysis.get("tables", [])
                    if tables:
                        meta.update(
                            {
                                "pdf_tables": {
                                    "table_count": len(tables),
                                    "total_rows": sum(t.get("shape", [0, 0])[0] for t in tables),
                                    "total_columns": sum(t.get("shape", [0, 0])[1] for t in tables),
                                    "extraction_methods": list(set(t.get("method", "unknown") for t in tables)),
                                }
                            }
                        )
                        log_info(f"PDF table extraction found {len(tables)} tables")

            except Exception as e:
                log_warning(f"Enhanced PDF processing failed: {e}")
                # Fallback to basic PDF processing
                try:
                    from pypdf import PdfReader

                    reader = PdfReader(str(file_path))
                    if not meta.get("title"):
                        meta["title"] = getattr(reader.metadata, "title", None)
                    if not meta.get("author"):
                        meta["author"] = getattr(reader.metadata, "author", None)
                except Exception:
                    pass

        elif file_type == ".docx":
            try:
                from docx import Document as DocxDocument

                doc = DocxDocument(str(file_path))
                props = doc.core_properties
                if not meta.get("author"):
                    meta["author"] = props.author
                if not meta.get("title"):
                    meta["title"] = props.title
            except Exception:
                pass

        elif file_type in {".xlsx", ".xls"}:
            try:
                import openpyxl

                wb = openpyxl.load_workbook(str(file_path), read_only=True)
                meta["sheets"] = wb.sheetnames
                meta["creator"] = wb.properties.creator
            except Exception:
                pass

        elif file_type in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}:
            # Enhanced image metadata already handled above
            pass

    except Exception as e:
        meta["meta_error"] = str(e)
        log_warning(f"Metadata extraction failed for {file_path}: {e}")

    if extra:
        meta.update(extra)
    return meta


def get_loader_for_file(file_path: Path):
    """
    Returns the appropriate document loader for a given file type.

    Args:
        file_path (Path): Path to the file.

    Returns:
        Document loader instance or None if unsupported.
    """
    suffix = file_path.suffix.lower()

    # Document formats
    if suffix == ".pdf":
        return PyPDFLoader(str(file_path))
    if suffix == ".docx":
        return Docx2txtLoader(str(file_path))
    if suffix in {".txt", ".md", ".rst", ".org", ".log", ".conf", ".properties"}:
        return TextLoader(str(file_path), autodetect_encoding=True)

    # Code files - treat as text
    if suffix in {
        ".py",
        ".js",
        ".ts",
        ".java",
        ".c",
        ".cpp",
        ".h",
        ".hpp",
        ".cs",
        ".go",
        ".rs",
        ".php",
        ".html",
        ".css",
        ".xml",
        ".json",
        ".yaml",
        ".yml",
        ".toml",
        ".ini",
        ".cfg",
        ".sql",
        ".sh",
        ".bat",
        ".ps1",
        ".r",
        ".m",
        ".swift",
        ".kt",
        ".dart",
    }:
        return TextLoader(str(file_path), autodetect_encoding=True)

    # Excel files - handled separately
    if suffix in {".xlsx", ".xls"}:
        return None  # Special handling in load_documents

    # Images - metadata only for now (future: OCR/vision)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}:
        return "metadata_only"

    # Email formats - need special handling
    if suffix in {".eml", ".msg", ".mbox"}:
        return "email_loader"  # Will implement this

    # Archives - need extraction
    if suffix in {".zip", ".tar", ".gz"}:
        return "archive_loader"  # Will implement this

    return None


def load_documents(folder_path: str) -> List[Document]:
    """
    Loads documents from the given folder with improved error handling and progress tracking.
    Supported formats: PDF, DOCX, Excel, TXT.

    Args:
        folder_path: Path to folder containing documents

    Returns:
        List of loaded documents
    """
    try:
        folder = validate_path(folder_path, must_exist=True)
    except (FileNotFoundError, NotADirectoryError) as e:
        log_error(str(e))
        return []

    supported_files = [f for f in folder.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not supported_files:
        log_warning(f"No supported files found in '{folder_path}'")
        return []

    docs, failed_files, ingest_events = [], [], []

    with create_progress_bar("Loading documents") as progress:
        task = progress.add_task("Loading...", total=len(supported_files))
        for file_path in supported_files:
            suffix = file_path.suffix.lower()
            try:
                meta = extract_metadata(file_path, suffix)
                if suffix in {".xlsx", ".xls"}:
                    file_docs = load_excel_with_pandas(file_path)
                    for doc in file_docs:
                        doc.metadata.update(meta)
                    docs.extend(file_docs)
                    ingest_events.append({"file": file_path.name, "status": "success", "meta": meta})
                else:
                    loader = get_loader_for_file(file_path)
                    if loader == "metadata_only":
                        # For images, check if we have OCR text
                        ocr_text = meta.get("ocr_text", "")
                        if ocr_text and len(ocr_text.strip()) > 10:
                            # Create document with OCR text as content
                            docs.append(Document(page_content=ocr_text, metadata=meta))
                            ingest_events.append(
                                {
                                    "file": file_path.name,
                                    "status": "success_with_ocr",
                                    "meta": meta,
                                    "ocr_length": len(ocr_text),
                                }
                            )
                        else:
                            # No useful text, just metadata
                            docs.append(Document(page_content="", metadata=meta))
                            ingest_events.append(
                                {
                                    "file": file_path.name,
                                    "status": "metadata_only",
                                    "meta": meta,
                                }
                            )
                    elif loader == "email_loader":
                        # Skip email files for now - not implemented
                        ingest_events.append(
                            {
                                "file": file_path.name,
                                "status": "skipped_email",
                                "meta": meta,
                            }
                        )
                    elif loader == "archive_loader":
                        # Skip archive files for now - not implemented
                        ingest_events.append(
                            {
                                "file": file_path.name,
                                "status": "skipped_archive",
                                "meta": meta,
                            }
                        )
                    elif loader:
                        file_docs = loader.load()
                        for doc in file_docs:
                            doc.metadata.update(meta)
                        docs.extend(file_docs)
                        ingest_events.append({"file": file_path.name, "status": "success", "meta": meta})
                    else:
                        ingest_events.append(
                            {
                                "file": file_path.name,
                                "status": "no_loader",
                                "meta": meta,
                            }
                        )
            except Exception as e:
                failed_files.append(file_path.name)
                ingest_events.append({"file": file_path.name, "status": "failed", "error": str(e)})
            progress.advance(task)

    log_dir = Path("data/ingest")
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"{datetime.now().strftime('%m-%d-%Y_%H-%M-%S')}.log"
    with open(log_file, "w", encoding="utf-8") as f:
        for event in ingest_events:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")

    if failed_files:
        log_warning(f"Failed to load {len(failed_files)} files: {', '.join(failed_files)}")
    log_success(f"Loaded {len(docs)} document chunks from {len(supported_files) - len(failed_files)} files")
    return docs


# Async versions with performance optimizations
@timer
async def rebuild_vectorstore_async(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
) -> Optional[FAISS]:
    """
    Async version: Wipes and rebuilds the FAISS vectorstore with parallel processing.

    This provides significant performance improvements through:
    - Parallel document loading
    - Cached embeddings (if previously computed)
    - Async I/O operations

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.

    Returns:
        Optional[FAISS]: The rebuilt FAISS vectorstore, or None if failed.
    """
    log_info("Starting async vectorstore rebuild...")

    # Initialize processor variable for cleanup
    processor = None

    try:
        # Use async document processor for parallel loading
        processor = AsyncDocumentProcessor()
        # Get all supported files
        folder = Path(folder_path)
        supported_files = [str(f) for f in folder.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS]

        if not supported_files:
            log_warning(f"No supported files found in '{folder_path}'")
            return None

        # Load documents in parallel (major performance boost!)
        log_info(f"📄 Loading {len(supported_files)} documents in parallel...")
        raw_docs = await processor.load_documents_from_paths(supported_files)

        if not raw_docs:
            log_error("No documents were loaded successfully")
            return None

        # Convert to expected format
        docs = []
        manifest = {"files": {}}

        for doc in raw_docs:
            docs.append({"text": doc.page_content, "metadata": doc.metadata})
            source_path = doc.metadata.get("source", "unknown")
            if Path(source_path).exists():
                manifest["files"][source_path] = file_checksum(Path(source_path))

        # Cleanup processor resources
        await processor.cleanup()
        processor = None  # Mark as cleaned up

        # Chunk documents (this is CPU-intensive but fast)
        log_info("🔪 Chunking documents...")
        chunks = chunk_documents(docs, chunk_size, chunk_overlap)
        texts = [c["text"] for c in chunks]
        metadatas = [c["metadata"] for c in chunks]

        # Build vectorstore with embeddings
        log_info("🧠 Creating embeddings and building vectorstore...")
        embeddings = OllamaEmbeddings(model=embedding_model)
        vectorstore = FAISS.from_texts(texts, embedding=embeddings, metadatas=metadatas)

        # Save to disk
        vectorstore.save_local(persist_path)
        save_manifest(persist_path, manifest)

        write_index_meta(
            persist_path,
            embedding_model=embedding_model,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        log_success(f"Async rebuild complete! {len(chunks)} chunks from {len(docs)} documents.")
        return vectorstore

    except Exception as e:
        log_error(f"Error during async rebuild: {e}")
        return None
    finally:
        # Ensure cleanup happens even if there's an error
        if processor is not None:
            try:
                await processor.cleanup()
            except Exception:
                pass


@timer
async def incremental_vectorstore_async(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
) -> Optional[FAISS]:
    """
    Async version: Incrementally updates vectorstore with only changed documents.

    Uses the incremental update system for maximum performance - only processes
    files that have actually changed since the last update.

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.

    Returns:
        Optional[FAISS]: The updated FAISS vectorstore, or None if failed.
    """
    log_info("🔄 Starting async incremental vectorstore update...")

    from core.processing.incremental_updates import IncrementalUpdateManager

    # Initialize incremental update manager
    update_manager = IncrementalUpdateManager(folder_path)

    # Process only changed documents
    log_info("🔍 Scanning for document changes...")
    update_results = await update_manager.process_incremental_update()

    if update_results.get("status") == "no_changes":
        log_info("✅ No changes detected - vectorstore is up to date!")
        # Load existing vectorstore
        from langchain_community.vectorstores import FAISS
        from langchain_ollama import OllamaEmbeddings

        try:
            embeddings = OllamaEmbeddings(model=embedding_model)
            vectorstore = FAISS.load_local(persist_path, embeddings, allow_dangerous_deserialization=True)
            return vectorstore
        except Exception as e:
            log_warning(f"Could not load existing vectorstore: {e}")
            # Fall back to full rebuild
            return await rebuild_vectorstore_async(folder_path, persist_path, chunk_size, chunk_overlap, embedding_model)

    elif update_results.get("status") == "success":
        # Process changed documents
        results = update_results.get("results", {})
        changed_files = results.get("added", []) + results.get("updated", [])

        if changed_files:
            log_info(f"📝 Processing {len(changed_files)} changed documents...")

            # Initialize processor variable for cleanup
            processor = None
            try:
                # Load only changed documents with async processing
                processor = AsyncDocumentProcessor()
                new_docs = await processor.load_documents_from_paths(changed_files)

                # Load existing vectorstore or create new one
                try:
                    embeddings = OllamaEmbeddings(model=embedding_model)
                    vectorstore = FAISS.load_local(persist_path, embeddings, allow_dangerous_deserialization=True)
                    log_info("📂 Loaded existing vectorstore")
                except Exception as e:
                    log_warning(f"Could not load existing vectorstore: {e}")
                    log_info("🆕 Creating new vectorstore...")
                    # Fall back to full rebuild
                    return await rebuild_vectorstore_async(
                        folder_path,
                        persist_path,
                        chunk_size,
                        chunk_overlap,
                        embedding_model,
                    )

                # Process new documents
                if new_docs:
                    # Convert and chunk new documents
                    docs = []
                    for doc in new_docs:
                        docs.append({"text": doc.page_content, "metadata": doc.metadata})

                    chunks = chunk_documents(docs, chunk_size, chunk_overlap)
                    texts = [c["text"] for c in chunks]
                    metadatas = [c["metadata"] for c in chunks]

                    # Add to existing vectorstore
                    log_info("➕ Adding new chunks to vectorstore...")
                    vectorstore.add_texts(texts, metadatas=metadatas)

                    # Save updated vectorstore
                    vectorstore.save_local(persist_path)

                    # Update manifest
                    manifest = load_manifest(persist_path)
                    for file_path in changed_files:
                        if Path(file_path).exists():
                            manifest["files"][file_path] = file_checksum(Path(file_path))
                    save_manifest(persist_path, manifest)

                    log_success(f"🎉 Incremental update complete! Added {len(chunks)} new chunks.")

                return vectorstore

            except Exception as e:
                log_error(f"Error during incremental processing: {e}")
                return None
            finally:
                # Cleanup processor
                if processor is not None:
                    try:
                        await processor.cleanup()
                    except Exception:
                        pass

    else:
        log_error("❌ Incremental update failed, falling back to full rebuild")
        return await rebuild_vectorstore_async(folder_path, persist_path, chunk_size, chunk_overlap, embedding_model)


@timer
def build_vectorstore(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
    mode: str = "rebuild",  # or "increment"
):
    """
    Builds or updates a FAISS vectorstore from documents in a folder.

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.
        mode (str): 'rebuild' to wipe and rebuild, 'increment' to update existing index.

    Returns:
        FAISS vectorstore or None if failed.
    """
    if mode == "rebuild":
        return rebuild_vectorstore(folder_path, persist_path, chunk_size, chunk_overlap, embedding_model)
    elif mode == "increment":
        return incremental_vectorstore(folder_path, persist_path, chunk_size, chunk_overlap, embedding_model)
    else:
        log_error(f"Invalid mode '{mode}'. Use 'rebuild' or 'increment'.")
        return None


@timer
def rebuild_vectorstore(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
) -> Optional[FAISS]:
    """
    Wipes and rebuilds the FAISS vectorstore from all documents in the folder.

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.

    Returns:
        Optional[FAISS]: The rebuilt FAISS vectorstore, or None if failed.
    """
    folder = Path(folder_path)
    manifest = {"files": {}}
    docs = []

    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        loader = get_loader_for_file(path)
        if loader == "metadata_only":
            text = ""
        elif loader == "email_loader" or loader == "archive_loader":
            continue  # Skip unimplemented loaders
        elif loader:
            text = "\n".join([d.page_content for d in loader.load()])
        else:
            continue
        if text.strip() or loader == "metadata_only":
            docs.append({"text": text, "metadata": {"source": str(path)}})
            manifest["files"][str(path)] = file_checksum(path)

    chunks = chunk_documents(docs, chunk_size, chunk_overlap)
    texts = [c["text"] for c in chunks]
    metadatas = [c["metadata"] for c in chunks]

    embeddings = OllamaEmbeddings(model=embedding_model)
    vectorstore = FAISS.from_texts(texts, embedding=embeddings, metadatas=metadatas)
    vectorstore.save_local(persist_path)
    save_manifest(persist_path, manifest)

    log_success(f"Rebuilt vectorstore with {len(chunks)} chunks from {len(docs)} documents.")
    return True


@timer
def incremental_vectorstore(
    folder_path: str,
    persist_path: str = "vectorstore",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = EMBEDDING_MODEL,
) -> Optional[FAISS]:
    """
    Incrementally updates an existing FAISS index:
    - Adds new/changed documents
    - Removes deleted documents

    Args:
        folder_path (str): Path to folder containing documents.
        persist_path (str): Path to save vectorstore.
        chunk_size (int): Size of text chunks for splitting.
        chunk_overlap (int): Overlap between chunks.
        embedding_model (str): Ollama embedding model to use.

    Returns:
        Optional[FAISS]: The updated FAISS vectorstore, or None if failed.
    """
    folder = Path(folder_path)
    old_manifest = load_manifest(persist_path)
    new_manifest = {"files": {}}
    changed_docs = []

    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        checksum = file_checksum(path)
        new_manifest["files"][str(path)] = checksum
        if old_manifest["files"].get(str(path)) != checksum:
            loader = get_loader_for_file(path)
            if loader == "metadata_only":
                text = ""
            elif loader == "email_loader" or loader == "archive_loader":
                continue  # Skip unimplemented loaders
            elif loader:
                text = "\n".join([d.page_content for d in loader.load()])
            else:
                continue
            if text.strip() or loader == "metadata_only":
                changed_docs.append({"text": text, "metadata": {"source": str(path)}})

    removed = set(old_manifest["files"]) - set(new_manifest["files"])
    if removed:
        log_info(f"Detected {len(removed)} deleted files: {removed}")

    if changed_docs or removed:
        return rebuild_vectorstore(folder_path, persist_path, chunk_size, chunk_overlap, embedding_model)

    log_info("No document changes detected. Vectorstore is up-to-date.")
    return True
